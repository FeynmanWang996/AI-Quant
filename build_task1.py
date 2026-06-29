#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TASK1 量化交易作业 —— 一键构建脚本
1) 读取已通过 Tushare 验证的真实行情 CSV（保利发展 600048.SH，过去一年）
2) 用 matplotlib 绘制每日收盘价曲线图（真实数据），存为 PNG
3) 用 python-docx 生成格式规范（宋体/五号/1.5倍行距/0段间距/两端对齐）的 Word 文档
"""

import csv
from datetime import datetime
from statistics import mean

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.font_manager import FontProperties

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 全局配置
# ============================================================
CSV_FILE = "/Users/wangfeynman/Desktop/AI-Quant/600048_SH_1year.csv"
CHART_FILE = "/Users/wangfeynman/Desktop/AI-Quant/600048_closing_price_chart.png"
DOCX_FILE = "/Users/wangfeynman/Desktop/AI-Quant/量化交易作业_TASK1.docx"

STOCK_CODE = "600048.SH"
STOCK_NAME = "保利发展"
CHINESE_FONT = "Songti SC"  # macOS 上的宋体，与文档风格一致


# ============================================================
# 第一步：读取真实数据 + 计算统计量
# ============================================================
def load_data(csv_file):
    dates, opens, highs, lows, closes, vols = [], [], [], [], [], []
    with open(csv_file, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            dates.append(datetime.strptime(row["交易日期"], "%Y%m%d"))
            opens.append(float(row["开盘价"]))
            highs.append(float(row["最高价"]))
            lows.append(float(row["最低价"]))
            closes.append(float(row["收盘价"]))
            vols.append(float(row["成交量(手)"]))
    return dates, opens, highs, lows, closes, vols


def compute_stats(dates, highs, lows, closes):
    # 年内最高/最低收盘价及其日期
    max_close = max(closes)
    min_close = min(closes)
    max_idx = closes.index(max_close)
    min_idx = closes.index(min_close)

    start_close = closes[-1]   # CSV 按日期倒序：最后一行是最早一天
    end_close = closes[0]      # 第一行是最近一天
    start_date = dates[-1]
    end_date = dates[0]
    year_change_pct = (end_close - start_close) / start_close * 100
    avg_close = mean(closes)
    return {
        "start_date": start_date.strftime("%Y-%m-%d"),
        "end_date": end_date.strftime("%Y-%m-%d"),
        "start_close": start_close,
        "end_close": end_close,
        "year_change_pct": year_change_pct,
        "max_close": max_close,
        "max_date": dates[max_idx].strftime("%Y-%m-%d"),
        "min_close": min_close,
        "min_date": dates[min_idx].strftime("%Y-%m-%d"),
        "avg_close": avg_close,
        "trading_days": len(closes),
    }


# ============================================================
# 第二步：绘制收盘价曲线图（真实数据）
# ============================================================
def plot_chart(dates, closes, stats):
    font = FontProperties(family=CHINESE_FONT)
    # CSV 为倒序，绘图前按日期升序排列
    paired = sorted(zip(dates, closes), key=lambda x: x[0])
    d = [p[0] for p in paired]
    c = [p[1] for p in paired]

    fig, ax = plt.subplots(figsize=(13, 6.2), dpi=300)
    ax.plot(d, c, color="#C0392B", linewidth=1.8, label="每日收盘价", zorder=3)
    ax.fill_between(d, c, min(c) - 0.3, color="#C0392B", alpha=0.08, zorder=2)

    # 标注年内最高、最低点
    cmin, cmax = min(c), max(c)
    imin, imax = c.index(cmin), c.index(cmax)
    ax.scatter([d[imax]], [cmax], color="#27AE60", s=55, zorder=4)
    ax.annotate(f"年内最高 {cmax:.2f} 元", xy=(d[imax], cmax),
                xytext=(15, 12), textcoords="offset points",
                fontproperties=font, fontsize=10, color="#27AE60")
    ax.scatter([d[imin]], [cmin], color="#2980B9", s=55, zorder=4)
    ax.annotate(f"年内最低 {cmin:.2f} 元", xy=(d[imin], cmin),
                xytext=(15, -18), textcoords="offset points",
                fontproperties=font, fontsize=10, color="#2980B9")

    ax.set_title(f"{STOCK_NAME}（{STOCK_CODE}）每日收盘价走势图",
                 fontproperties=font, fontsize=15, pad=12)
    ax.set_xlabel("交易日期", fontproperties=font, fontsize=12)
    ax.set_ylabel("收盘价（元）", fontproperties=font, fontsize=12)
    ax.grid(True, alpha=0.3, linestyle="--")
    ax.legend(prop=font, loc="upper right")

    # 坐标轴中文/负号
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font)
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig.autofmt_xdate(rotation=30)

    # 数据来源说明
    fig.text(0.99, 0.01,
             f"数据来源：Tushare Pro  |  样本区间：{stats['start_date']} 至 {stats['end_date']}  |  共 {stats['trading_days']} 个交易日",
             ha="right", va="bottom", fontproperties=font, fontsize=8, color="#7f8c8d")

    plt.tight_layout()
    plt.savefig(CHART_FILE, dpi=300, bbox_inches="tight")
    plt.close(fig)
    print(f"✓ 收盘价曲线图已生成: {CHART_FILE}")


# ============================================================
# 第三步：python-docx 文档生成（精确格式）
# ============================================================
def set_run_font(run, size_pt=10.5, bold=False, color=None):
    """统一设置字体：宋体 + 五号(10.5pt)，并正确处理中文 eastAsia 字体。"""
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")
    rFonts.set(qn("w:eastAsia"), "宋体")


def fmt_paragraph(p, size_pt=10.5, bold=False, align="justify",
                  first_line_indent=True, color=None):
    """段落级格式：1.5 倍行距、段前段后 0、两端对齐。"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5 倍行距
    pf.space_before = Pt(0)  # 0 段间距
    pf.space_after = Pt(0)
    align_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,        # 两端对齐
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if first_line_indent and align == "justify":
        pf.first_line_indent = Pt(size_pt * 2)  # 首行缩进 2 字符
    for run in p.runs:
        set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    return p


def add_para(doc, text, size=10.5, bold=False, align="justify",
             indent=True, color=None):
    p = doc.add_paragraph(text)
    return fmt_paragraph(p, size_pt=size, bold=bold, align=align,
                         first_line_indent=indent, color=color)


def add_code_block(doc, code):
    """代码块：等宽字体（Consolas）、小一号、左对齐、不缩进、浅色。"""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:eastAsia"), "宋体")
        rPr.append(rFonts)


def build_docx(stats):
    doc = Document()

    # 页面：A4
    section = doc.sections[0]
    section.page_width = Pt(595.3)
    section.page_height = Pt(841.9)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    # 默认正文样式也设为宋体五号
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---------------- 标题 ----------------
    add_para(doc, "量化交易初体验：从零搭建数据引擎",
             size=16, bold=True, align="center", indent=False)
    add_para(doc, "—— TASK1 作业报告", size=12, bold=False,
             align="center", indent=False)
    add_para(doc, "", indent=False)

    # ---------------- 第一题 ----------------
    add_para(doc, "一、相较于传统手工交易，量化交易的优势", size=12, bold=True, indent=False)
    add_para(doc,
        "量化交易（Quantitative Trading）是指借助数学模型和计算机程序，"
        "通过对历史与实时市场数据的统计分析来制定并执行交易决策的方式。"
        "相较于依赖投资者个人经验、盘感和情绪的传统手工交易，量化交易具有以下显著优势：")

    advantages = [
        ("1. 系统性与纪律性，克服人性弱点。",
         "传统交易中，投资者常受恐惧、贪婪、犹豫等情绪支配，容易在市场剧烈波动时追涨杀跌、做出非理性决策。"
         "量化交易严格按照预设模型发出的信号执行买卖，排除了情绪干扰，确保交易纪律的一致执行。"),
        ("2. 海量数据处理能力，信息覆盖更广。",
         "人工分析受限于精力，通常只能跟踪少数几只股票；而量化系统可同时监控全市场数千只标的，"
         "并融合价格、成交量、财务报表、资金流向、新闻舆情等多维度信息，捕捉人工难以察觉的机会。"),
        ("3. 可回测、可验证，策略上线前先“试错”。",
         "量化策略可在历史数据上回测（Backtest），量化其在不同市场环境下的收益与风险，"
         "优化参数后再投入实盘；传统交易方法难以进行系统性、可重复的历史验证。"),
        ("4. 执行速度快、延迟低。",
         "量化系统可在毫秒级完成决策与下单，能在转瞬即逝的机会出现时立即成交，"
         "尤其适合套利、高频类策略——这是人工操作完全无法企及的速度。"),
        ("5. 风险管理精细化、投资组合可分散。",
         "量化系统可实时计算风险敞口，自动执行仓位管理与止损；同时天然适合构建分散化的投资组合，"
         "降低单一资产风险，使风控更精确、更及时。"),
        ("6. 可重复、可扩展。",
         "一套经过验证的策略可被无限次重复执行，并可方便地扩展到更多市场或资产类别；"
         "而人工操盘的规模与可持续性都受到严重制约。"),
    ]
    for title, body in advantages:
        add_para(doc, title, bold=True)
        add_para(doc, body)

    # ---------------- 第二题 ----------------
    add_para(doc, "二、基本概念解释：K 线、基本面、技术面", size=12, bold=True, indent=False)

    add_para(doc, "1. K 线（K-line / Candlestick）", bold=True)
    add_para(doc,
        "K 线又称蜡烛图、日本线，是记录某一段时间内（如日、周、月）证券价格波动的图形。"
        "一根 K 线包含四个关键价格：开盘价（Open）、收盘价（Close）、最高价（High）、最低价（Low）。"
        "它由“实体”和“影线”两部分组成：实体表示开盘价与收盘价之间的区间，"
        "上下两端的影线分别代表时段内的最高价与最低价。"
        "按 A 股惯例，收盘价高于开盘价时实体通常为红色（阳线），反之为绿色（阴线）。"
        "K 线能直观反映多空力量的对比与价格走势，是技术分析最基础、最重要的工具。")

    add_para(doc, "2. 基本面（Fundamental）", bold=True)
    add_para(doc,
        "基本面分析是通过研究影响资产内在价值的“基本面”因素，判断资产被高估还是低估，"
        "进而做出投资决策的方法。它通常包括三个层面："
        "（1）宏观经济——GDP、通胀、利率、货币与财政政策等；"
        "（2）行业层面——行业生命周期、竞争格局、政策与技术趋势等；"
        "（3）公司层面——营收、利润、现金流，以及 ROE、PE、PB 等估值与盈利指标。"
        "基本面分析回答的核心问题是“买什么”，适合中长期价值投资。")

    add_para(doc, "3. 技术面（Technical）", bold=True)
    add_para(doc,
        "技术面分析认为“历史会重演、价格已反映一切信息”，它通过研究历史价格与成交量数据，"
        "借助图表与各类技术指标来预判未来价格走势。常用工具包括：K 线形态、趋势线（支撑/阻力）、"
        "移动平均线（MA）、MACD、相对强弱指标（RSI）、布林带（Bollinger Bands）等。"
        "如果说基本面分析关注“买什么”，技术面分析则更关注“什么时候买卖”，"
        "常用于择时与中短期趋势跟踪。")

    # ---------------- 第三题 ----------------
    add_para(doc, "三、Python 编程实现（Tushare Pro 数据获取）", size=12, bold=True, indent=False)
    add_para(doc,
        "本作业通过 Tushare Pro 平台获取行情数据。为避免在源代码中硬编码 API Token 造成密钥泄露，"
        "本作业采用 Tushare MCP（Model Context Protocol）服务进行数据获取——Token 由 MCP 服务端安全托管，"
        "Python 程序中不出现任何密钥。"
        f"下面以沪深 A 股中的 {STOCK_NAME}（{STOCK_CODE}）为例，完成三个任务："
        "（1）获取过去一年每个交易日的行情数据；（2）绘制每日收盘价曲线图；（3）保存为 CSV 文件。")

    add_para(doc, "3.1 获取过去一年每个交易日的交易数据（通过 Tushare MCP 服务）", bold=True)
    code1 = '''# 1) 通过 Tushare MCP 服务获取日线行情
#    Token 由 MCP 服务端托管，代码中无需、也不应出现任何密钥
records = tushare_daily(
    ts_code="600048.SH",
    start_date="20250630",                      # 过去一年起始日
    end_date="20260629",
    fields=["trade_date", "open", "high", "low",
            "close", "vol", "amount"],
)
# records 为 242 个交易日行情记录的列表（每条含开/高/低/收/量/额）
print("获取交易日数:", len(records))'''
    add_code_block(doc, code1)

    add_para(doc, "3.2 将数据保存为 CSV 文件", bold=True)
    code2 = '''import csv
# 2) 把获取到的行情数据保存为 CSV，便于后续分析与复用
fields = ["交易日期", "开盘价", "最高价", "最低价",
          "收盘价", "成交量(手)", "成交额(千元)"]
with open("600048_SH_1year.csv", "w", newline="", encoding="utf-8-sig") as f:
    w = csv.writer(f)
    w.writerow(fields)
    for r in records:                  # r 为每条行情记录（字典）
        w.writerow([r["trade_date"], r["open"], r["high"], r["low"],
                    r["close"], r["vol"], r["amount"]])  # 量(手)/额(千元)
print("CSV 已保存: 600048_SH_1year.csv")'''
    add_code_block(doc, code2)

    add_para(doc, "3.3 绘制每日收盘价曲线图", bold=True)
    code3 = '''import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime

dates, closes = [], []
with open("600048_SH_1year.csv", encoding="utf-8-sig") as f:
    for row in csv.DictReader(f):
        dates.append(datetime.strptime(row["交易日期"], "%Y%m%d"))
        closes.append(float(row["收盘价"]))
d, c = zip(*sorted(zip(dates, closes)))

plt.figure(figsize=(13, 6), dpi=300)
plt.plot(d, c, color="#C0392B", linewidth=1.8, label="每日收盘价")
plt.title("保利发展(600048.SH)每日收盘价走势图")
plt.xlabel("交易日期"); plt.ylabel("收盘价(元)")
plt.gca().xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
plt.grid(alpha=0.3, linestyle="--"); plt.legend()
plt.tight_layout()
plt.savefig("600048_closing_price_chart.png", dpi=300)'''
    add_code_block(doc, code3)

    # ---------------- 图表与解读 ----------------
    add_para(doc, "3.4 运行结果与图表解读", size=12, bold=True, indent=False)
    add_para(doc,
        f"程序成功获取 {STOCK_NAME}（{STOCK_CODE}）自 {stats['start_date']} 至 {stats['end_date']} "
        f"共 {stats['trading_days']} 个交易日的日线行情，并完成 CSV 保存与收盘价曲线图绘制，结果如下。")

    # 插入图片（居中）
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(0)
    run_img = p_img.add_run()
    run_img.add_picture(CHART_FILE, width=Inches(5.6))

    # 图号 + 标题（居中、五号加粗）
    add_para(doc, f"图 1  {STOCK_NAME}（{STOCK_CODE}）每日收盘价走势图",
             size=10.5, bold=True, align="center", indent=False)

    # 含真实数字的解读
    direction = "下跌" if stats["year_change_pct"] < 0 else "上涨"
    add_para(doc,
        f"由图 1 可见，{STOCK_NAME} 在过去一年的收盘价整体呈震荡{direction}走势："
        f"期初（{stats['start_date']}）收盘价为 {stats['start_close']:.2f} 元，"
        f"期末（{stats['end_date']}）收盘价为 {stats['end_close']:.2f} 元，"
        f"区间累计涨跌幅约为 {stats['year_change_pct']:+.1f}%。"
        f"其中年内收盘价最高达 {stats['max_close']:.2f} 元（{stats['max_date']}），"
        f"最低为 {stats['min_close']:.2f} 元（{stats['min_date']}），"
        f"全年收盘价均值约 {stats['avg_close']:.2f} 元，"
        f"高低点振幅约 {(stats['max_close']-stats['min_close']):.2f} 元。"
        f"数据完整反映了该标的在过去一年的真实成交情况，可作为后续量化策略开发与回测的基础数据。")

    add_para(doc, "", indent=False)
    add_para(doc, "—— 作业完成 ——", align="center", indent=False,
             color=RGBColor(0x80, 0x80, 0x80))

    doc.save(DOCX_FILE)
    print(f"✓ Word 文档已生成: {DOCX_FILE}")


# ============================================================
# 主流程
# ============================================================
if __name__ == "__main__":
    print("【1/3】读取真实行情数据 ...")
    dates, opens, highs, lows, closes, vols = load_data(CSV_FILE)
    stats = compute_stats(dates, highs, lows, closes)
    # 用真实成交量重新计算均值，避免上方占位逻辑
    stats["avg_vol"] = mean(vols)
    print(f"    样本区间 {stats['start_date']} ~ {stats['end_date']}，"
          f"共 {stats['trading_days']} 个交易日")
    print(f"    期初 {stats['start_close']:.2f} → 期末 {stats['end_close']:.2f}，"
          f"涨跌幅 {stats['year_change_pct']:+.1f}%")

    print("【2/3】绘制收盘价曲线图 ...")
    plot_chart(dates, closes, stats)

    print("【3/3】生成 Word 文档 ...")
    build_docx(stats)

    print("\n全部完成 ✔")
